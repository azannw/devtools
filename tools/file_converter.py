"""
ToolKit — File Type Converter.
Convert between document formats and image formats.
"""

import csv
import os
import threading
from pathlib import Path
from tkinter import filedialog

import customtkinter as ctk
from PIL import Image

from utils import theme


# ── Conversion registry ──────────────────────────────────────────────────

# Maps (src_ext, dst_ext) → handler(input_path, output_path, **kwargs)
_DOC_CONVERTERS = {}
_DOC_TARGETS = {}  # src_ext → [possible target exts]


def _register_doc(src, dst, func):
    _DOC_CONVERTERS[(src, dst)] = func
    _DOC_TARGETS.setdefault(src, []).append(dst)


# ── Document conversion functions ─────────────────────────────────────────

def _docx_to_pdf(inp, out, **kw):
    try:
        import comtypes.client
        word = comtypes.client.CreateObject("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(str(Path(inp).resolve()))
        doc.SaveAs(str(Path(out).resolve()), FileFormat=17)
        doc.Close()
        word.Quit()
    except Exception:
        # fallback: basic conversion via python-docx + fpdf2
        from docx import Document
        from fpdf import FPDF
        d = Document(inp)
        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.set_font("Helvetica", size=11)
        for para in d.paragraphs:
            text = para.text
            if text.strip():
                pdf.multi_cell(0, 6, text)
                pdf.ln(2)
        pdf.output(out)


def _docx_to_txt(inp, out, **kw):
    from docx import Document
    doc = Document(inp)
    text = "\n".join(p.text for p in doc.paragraphs)
    Path(out).write_text(text, encoding="utf-8")


def _txt_to_pdf(inp, out, **kw):
    from fpdf import FPDF
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_font("Helvetica", size=11)
    text = Path(inp).read_text(encoding="utf-8", errors="replace")
    for line in text.split("\n"):
        pdf.multi_cell(0, 6, line)
        pdf.ln(1)
    pdf.output(out)


def _txt_to_docx(inp, out, **kw):
    from docx import Document
    doc = Document()
    text = Path(inp).read_text(encoding="utf-8", errors="replace")
    for line in text.split("\n"):
        doc.add_paragraph(line)
    doc.save(out)


def _pdf_to_txt(inp, out, **kw):
    from pypdf import PdfReader
    reader = PdfReader(inp)
    text = "\n".join(page.extract_text() or "" for page in reader.pages)
    Path(out).write_text(text, encoding="utf-8")


def _pdf_to_docx(inp, out, **kw):
    try:
        from pdf2docx import Converter
        cv = Converter(inp)
        cv.convert(out, start=0, end=None)
        cv.close()
    except Exception:
        # fallback: extract text and dump into docx
        from pypdf import PdfReader
        from docx import Document
        reader = PdfReader(inp)
        doc = Document()
        for page in reader.pages:
            text = page.extract_text() or ""
            for line in text.split("\n"):
                doc.add_paragraph(line)
        doc.save(out)


def _pptx_to_pdf(inp, out, **kw):
    try:
        import comtypes.client
        pp = comtypes.client.CreateObject("PowerPoint.Application")
        pres = pp.Presentations.Open(str(Path(inp).resolve()), WithWindow=False)
        pres.SaveAs(str(Path(out).resolve()), 32)  # ppSaveAsPDF
        pres.Close()
        pp.Quit()
    except Exception as e:
        raise RuntimeError(f"PowerPoint conversion requires Microsoft Office. {e}")


def _xlsx_to_pdf(inp, out, **kw):
    try:
        import comtypes.client
        excel = comtypes.client.CreateObject("Excel.Application")
        excel.Visible = False
        wb = excel.Workbooks.Open(str(Path(inp).resolve()))
        wb.ExportAsFixedFormat(0, str(Path(out).resolve()))
        wb.Close(False)
        excel.Quit()
    except Exception as e:
        raise RuntimeError(f"Excel conversion requires Microsoft Office. {e}")


def _xlsx_to_csv(inp, out, **kw):
    from openpyxl import load_workbook
    wb = load_workbook(inp, read_only=True)
    ws = wb.active
    with open(out, "w", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        for row in ws.iter_rows(values_only=True):
            writer.writerow(row)
    wb.close()


def _csv_to_xlsx(inp, out, **kw):
    from openpyxl import Workbook
    wb = Workbook()
    ws = wb.active
    with open(inp, "r", encoding="utf-8", errors="replace") as f:
        reader = csv.reader(f)
        for row in reader:
            ws.append(row)
    wb.save(out)


def _html_to_pdf(inp, out, **kw):
    from fpdf import FPDF
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_font("Helvetica", size=11)
    html = Path(inp).read_text(encoding="utf-8", errors="replace")
    # Strip to just text with basic HTML support
    try:
        pdf.write_html(html)
    except Exception:
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(html, "html.parser")
        pdf.multi_cell(0, 6, soup.get_text())
    pdf.output(out)


def _md_to_pdf(inp, out, **kw):
    from markdown import markdown
    from fpdf import FPDF
    md_text = Path(inp).read_text(encoding="utf-8", errors="replace")
    html = markdown(md_text, extensions=["tables", "fenced_code"])
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_font("Helvetica", size=11)
    try:
        pdf.write_html(html)
    except Exception:
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(html, "html.parser")
        pdf.multi_cell(0, 6, soup.get_text())
    pdf.output(out)


# register all document conversions
_register_doc("docx", "pdf",  _docx_to_pdf)
_register_doc("docx", "txt",  _docx_to_txt)
_register_doc("txt",  "pdf",  _txt_to_pdf)
_register_doc("txt",  "docx", _txt_to_docx)
_register_doc("pdf",  "txt",  _pdf_to_txt)
_register_doc("pdf",  "docx", _pdf_to_docx)
_register_doc("pptx", "pdf",  _pptx_to_pdf)
_register_doc("xlsx", "pdf",  _xlsx_to_pdf)
_register_doc("xlsx", "csv",  _xlsx_to_csv)
_register_doc("csv",  "xlsx", _csv_to_xlsx)
_register_doc("html", "pdf",  _html_to_pdf)
_register_doc("md",   "pdf",  _md_to_pdf)

# supported image extensions
_IMAGE_EXTS = ["png", "jpg", "jpeg", "webp", "bmp", "ico", "tiff", "gif"]
_IMAGE_OUTPUT_FORMATS = ["PNG", "JPEG", "WebP", "BMP", "ICO", "TIFF", "GIF"]


# ── Tool UI ──────────────────────────────────────────────────────────────

class FileConverterTool(ctk.CTkFrame):
    def __init__(self, parent):
        super().__init__(parent, fg_color=theme.BG_DARK)
        self._doc_path = None
        self._image_paths = []
        self._category = "Documents"
        self._build_ui()

    def _build_ui(self):
        self.grid_columnconfigure(0, weight=1)
        self.grid_rowconfigure(2, weight=1)

        # header
        header = theme.create_tool_header(
            self, "File Converter",
            "Convert between document and image formats",
        )
        header.grid(row=0, column=0, sticky="ew",
                    padx=theme.PAD_LARGE, pady=(theme.PAD_LARGE, theme.PAD_MEDIUM))

        # category selector
        self._cat_seg = ctk.CTkSegmentedButton(
            self, values=["Documents", "Images"], command=self._switch_category,
            font=theme.FONT_BODY,
            fg_color=theme.BG_INPUT, selected_color=theme.ACCENT,
            selected_hover_color=theme.ACCENT_HOVER,
            unselected_color=theme.BG_INPUT, unselected_hover_color=theme.BG_CARD_HOVER,
            text_color=theme.TEXT_PRIMARY,
        )
        self._cat_seg.set("Documents")
        self._cat_seg.grid(row=1, column=0, sticky="w",
                          padx=theme.PAD_LARGE, pady=(0, theme.PAD_MEDIUM))

        # panels container
        self._panels = ctk.CTkFrame(self, fg_color="transparent")
        self._panels.grid(row=2, column=0, sticky="nsew", padx=theme.PAD_LARGE)
        self._panels.grid_columnconfigure(0, weight=1)
        self._panels.grid_rowconfigure(0, weight=1)

        self._doc_panel = self._build_doc_panel(self._panels)
        self._img_panel = self._build_img_panel(self._panels)
        self._doc_panel.grid(row=0, column=0, sticky="nsew")

        # status bar
        sf, self._status_label, self._progress_bar = theme.create_status_bar(self)
        sf.grid(row=3, column=0, sticky="ew",
                padx=theme.PAD_LARGE, pady=(theme.PAD_MEDIUM, theme.PAD_LARGE))

    def _switch_category(self, value):
        self._category = value
        if value == "Documents":
            self._img_panel.grid_forget()
            self._doc_panel.grid(row=0, column=0, sticky="nsew")
        else:
            self._doc_panel.grid_forget()
            self._img_panel.grid(row=0, column=0, sticky="nsew")

    # ── documents panel ───────────────────────────────────────────────

    def _build_doc_panel(self, parent):
        card = theme.create_card_frame(parent)
        inner = ctk.CTkFrame(card, fg_color="transparent")
        inner.pack(fill="both", expand=True, padx=theme.PAD_LARGE, pady=theme.PAD_LARGE)

        # select file
        top = ctk.CTkFrame(inner, fg_color="transparent")
        top.pack(fill="x")
        theme.create_action_button(top, "Select File", self._select_doc, width=180).pack(side="left")
        self._doc_label = ctk.CTkLabel(top, text="No file selected", font=theme.FONT_SMALL, text_color=theme.TEXT_SECONDARY)
        self._doc_label.pack(side="left", padx=theme.PAD_MEDIUM)

        # supported formats note
        doc_exts = sorted(set(_DOC_TARGETS.keys()))
        ctk.CTkLabel(
            inner,
            text=f"Supported: {', '.join('.' + e for e in doc_exts)}",
            font=theme.FONT_SMALL, text_color=theme.TEXT_DISABLED,
        ).pack(anchor="w", pady=(theme.PAD_SMALL, theme.PAD_MEDIUM))

        # output format
        fmt_row = ctk.CTkFrame(inner, fg_color="transparent")
        fmt_row.pack(fill="x")
        ctk.CTkLabel(fmt_row, text="Convert to", font=theme.FONT_SMALL, text_color=theme.TEXT_SECONDARY).pack(side="left")
        self._doc_fmt_menu = ctk.CTkOptionMenu(
            fmt_row, values=["Select a file first"],
            fg_color=theme.BG_INPUT, button_color=theme.ACCENT,
            button_hover_color=theme.ACCENT_HOVER, text_color=theme.TEXT_PRIMARY,
            dropdown_fg_color=theme.BG_CARD, dropdown_hover_color=theme.ACCENT_MUTED,
            dropdown_text_color=theme.TEXT_PRIMARY, width=200,
        )
        self._doc_fmt_menu.pack(side="right")

        # convert
        self._doc_convert_btn = theme.create_action_button(inner, "Convert", self._start_doc_convert, width=200)
        self._doc_convert_btn.pack(anchor="w", pady=(theme.PAD_LARGE, 0))

        # result
        self._doc_result_label = ctk.CTkLabel(inner, text="", font=theme.FONT_SMALL, text_color=theme.SUCCESS)
        self._doc_result_label.pack(anchor="w", pady=(theme.PAD_SMALL, 0))

        return card

    # ── images panel ──────────────────────────────────────────────────

    def _build_img_panel(self, parent):
        card = theme.create_card_frame(parent)
        inner = ctk.CTkFrame(card, fg_color="transparent")
        inner.pack(fill="both", expand=True, padx=theme.PAD_LARGE, pady=theme.PAD_LARGE)

        # select images
        top = ctk.CTkFrame(inner, fg_color="transparent")
        top.pack(fill="x")
        theme.create_action_button(top, "Select Images", self._select_images, width=180).pack(side="left")
        self._img_count_label = ctk.CTkLabel(top, text="No files selected", font=theme.FONT_SMALL, text_color=theme.TEXT_SECONDARY)
        self._img_count_label.pack(side="left", padx=theme.PAD_MEDIUM)

        # output format
        fmt_row = ctk.CTkFrame(inner, fg_color="transparent")
        fmt_row.pack(fill="x", pady=(theme.PAD_MEDIUM, 0))
        ctk.CTkLabel(fmt_row, text="Convert to", font=theme.FONT_SMALL, text_color=theme.TEXT_SECONDARY).pack(side="left")
        self._img_fmt_menu = ctk.CTkOptionMenu(
            fmt_row, values=_IMAGE_OUTPUT_FORMATS,
            command=self._on_img_format_change,
            fg_color=theme.BG_INPUT, button_color=theme.ACCENT,
            button_hover_color=theme.ACCENT_HOVER, text_color=theme.TEXT_PRIMARY,
            dropdown_fg_color=theme.BG_CARD, dropdown_hover_color=theme.ACCENT_MUTED,
            dropdown_text_color=theme.TEXT_PRIMARY, width=160,
        )
        self._img_fmt_menu.set("PNG")
        self._img_fmt_menu.pack(side="right")

        # quality slider (conditionally visible)
        self._img_quality_frame, self._img_quality_slider, _ = theme.create_slider_with_label(
            inner, "Quality", 1, 100, 85,
        )

        # ICO size selector (conditionally visible)
        self._ico_frame = ctk.CTkFrame(inner, fg_color="transparent")
        ctk.CTkLabel(self._ico_frame, text="ICO size", font=theme.FONT_SMALL, text_color=theme.TEXT_SECONDARY).pack(side="left")
        self._ico_menu = ctk.CTkOptionMenu(
            self._ico_frame, values=["16x16", "32x32", "48x48", "64x64", "128x128", "256x256"],
            fg_color=theme.BG_INPUT, button_color=theme.ACCENT,
            button_hover_color=theme.ACCENT_HOVER, text_color=theme.TEXT_PRIMARY,
            dropdown_fg_color=theme.BG_CARD, dropdown_hover_color=theme.ACCENT_MUTED,
            dropdown_text_color=theme.TEXT_PRIMARY, width=120,
        )
        self._ico_menu.set("256x256")
        self._ico_menu.pack(side="right")

        self._on_img_format_change("PNG")  # initial visibility

        # output folder
        out_row = ctk.CTkFrame(inner, fg_color="transparent")
        out_row.pack(fill="x", pady=(theme.PAD_MEDIUM, 0))
        theme.create_secondary_button(out_row, "Output Folder", self._select_img_output, width=140).pack(side="left")
        self._img_out_label = ctk.CTkLabel(out_row, text="Same as input", font=theme.FONT_SMALL, text_color=theme.TEXT_SECONDARY)
        self._img_out_label.pack(side="left", padx=theme.PAD_SMALL)
        self._img_output_dir = None

        # convert
        self._img_convert_btn = theme.create_action_button(inner, "Convert All", self._start_img_convert, width=200)
        self._img_convert_btn.pack(anchor="w", pady=(theme.PAD_LARGE, 0))

        return card

    # ── callbacks ─────────────────────────────────────────────────────

    def _select_doc(self):
        ext_list = " ".join(f"*.{e}" for e in _DOC_TARGETS.keys())
        path = filedialog.askopenfilename(
            title="Select Document",
            filetypes=[("Documents", ext_list), ("All files", "*.*")],
        )
        if not path:
            return
        self._doc_path = path
        self._doc_label.configure(text=Path(path).name)
        self._doc_result_label.configure(text="")

        ext = Path(path).suffix.lstrip(".").lower()
        targets = _DOC_TARGETS.get(ext, [])
        if targets:
            display = [t.upper() for t in targets]
            self._doc_fmt_menu.configure(values=display)
            self._doc_fmt_menu.set(display[0])
        else:
            self._doc_fmt_menu.configure(values=["Unsupported format"])
            self._doc_fmt_menu.set("Unsupported format")

    def _select_images(self):
        ext_str = " ".join(f"*.{e}" for e in _IMAGE_EXTS)
        paths = filedialog.askopenfilenames(
            title="Select Images",
            filetypes=[("Images", ext_str)],
        )
        if paths:
            self._image_paths = list(paths)
            self._img_count_label.configure(text=f"{len(paths)} file(s) selected")

    def _select_img_output(self):
        d = filedialog.askdirectory(title="Select output folder")
        if d:
            self._img_output_dir = d
            self._img_out_label.configure(text=d)

    def _on_img_format_change(self, value):
        # show/hide quality slider
        self._img_quality_frame.pack_forget()
        self._ico_frame.pack_forget()
        if value in ("JPEG", "WebP"):
            self._img_quality_frame.pack(fill="x", pady=(theme.PAD_MEDIUM, 0))
        elif value == "ICO":
            self._ico_frame.pack(fill="x", pady=(theme.PAD_MEDIUM, 0))

    # ── document conversion ───────────────────────────────────────────

    def _start_doc_convert(self):
        if not self._doc_path:
            self._status_label.configure(text="Select a file first", text_color=theme.WARNING)
            return

        src_ext = Path(self._doc_path).suffix.lstrip(".").lower()
        dst_ext = self._doc_fmt_menu.get().lower()
        key = (src_ext, dst_ext)

        if key not in _DOC_CONVERTERS:
            self._status_label.configure(text=f"Conversion {src_ext} → {dst_ext} not supported", text_color=theme.ERROR)
            return

        out = filedialog.asksaveasfilename(
            title="Save converted file",
            defaultextension=f".{dst_ext}",
            filetypes=[(dst_ext.upper(), f"*.{dst_ext}")],
            initialfile=Path(self._doc_path).stem + f".{dst_ext}",
        )
        if not out:
            return

        self._doc_convert_btn.configure(state="disabled")
        self._status_label.configure(text="Converting…", text_color=theme.TEXT_SECONDARY)
        self._progress_bar.set(0)

        def _run():
            try:
                self.after(0, lambda: self._progress_bar.configure(mode="indeterminate"))
                self.after(0, lambda: self._progress_bar.start())
                _DOC_CONVERTERS[key](self._doc_path, out)
                self.after(0, lambda: self._progress_bar.stop())
                self.after(0, lambda: self._progress_bar.configure(mode="determinate"))
                self.after(0, lambda: self._progress_bar.set(1))
                self.after(0, lambda: self._status_label.configure(
                    text=f"Converted to {Path(out).name}", text_color=theme.SUCCESS))
                self.after(0, lambda: self._doc_result_label.configure(text=f"Saved: {out}"))
            except Exception as e:
                self.after(0, lambda: self._progress_bar.stop())
                self.after(0, lambda: self._progress_bar.configure(mode="determinate"))
                self.after(0, lambda: self._status_label.configure(text=f"Error: {e}", text_color=theme.ERROR))
            finally:
                self.after(0, lambda: self._doc_convert_btn.configure(state="normal"))

        threading.Thread(target=_run, daemon=True).start()

    # ── image conversion ──────────────────────────────────────────────

    def _start_img_convert(self):
        if not self._image_paths:
            self._status_label.configure(text="Select images first", text_color=theme.WARNING)
            return

        self._img_convert_btn.configure(state="disabled")
        self._status_label.configure(text="Converting…", text_color=theme.TEXT_SECONDARY)
        self._progress_bar.set(0)
        threading.Thread(target=self._do_img_convert, daemon=True).start()

    def _do_img_convert(self):
        try:
            fmt = self._img_fmt_menu.get()
            quality = int(self._img_quality_slider.get())
            total = len(self._image_paths)

            ext_map = {
                "PNG": ".png", "JPEG": ".jpg", "WebP": ".webp",
                "BMP": ".bmp", "ICO": ".ico", "TIFF": ".tiff", "GIF": ".gif",
            }
            out_ext = ext_map.get(fmt, ".png")

            for i, path in enumerate(self._image_paths):
                img = Image.open(path)
                out_dir = self._img_output_dir or str(Path(path).parent)
                out_path = os.path.join(out_dir, Path(path).stem + out_ext)

                # handle alpha for formats that don't support it
                if fmt in ("JPEG", "BMP") and img.mode in ("RGBA", "LA", "PA"):
                    bg = Image.new("RGB", img.size, (255, 255, 255))
                    bg.paste(img, mask=img.convert("RGBA").split()[3])
                    img = bg
                elif fmt in ("JPEG", "BMP") and img.mode != "RGB":
                    img = img.convert("RGB")

                if fmt == "ICO":
                    size_str = self._ico_menu.get()
                    s = int(size_str.split("x")[0])
                    img = img.resize((s, s), Image.LANCZOS)
                    img.save(out_path, format="ICO", sizes=[(s, s)])
                elif fmt in ("JPEG", "WebP"):
                    img.save(out_path, format=fmt, quality=quality, optimize=True)
                elif fmt == "GIF":
                    img.save(out_path, format="GIF")
                else:
                    img.save(out_path, format=fmt)

                prog = (i + 1) / total
                self.after(0, lambda p=prog: self._progress_bar.set(p))
                self.after(0, lambda ii=i + 1: self._status_label.configure(text=f"Converting {ii}/{total}…"))

            self.after(0, lambda: self._status_label.configure(
                text=f"Converted {total} image(s) to {fmt}", text_color=theme.SUCCESS))
            self.after(0, lambda: self._progress_bar.set(1))
        except Exception as e:
            self.after(0, lambda: self._status_label.configure(text=f"Error: {e}", text_color=theme.ERROR))
        finally:
            self.after(0, lambda: self._img_convert_btn.configure(state="normal"))
